"""Turn an uploaded file into plain text plus layout signals.

The layout signals matter as much as the text: multi-column layouts, text in
tables and image-only PDFs are the main reasons an applicant tracking system
mis-reads a resume, and none of that is visible in the extracted string.
"""

from __future__ import annotations

import hashlib
import io
import logging
import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import PurePath

logger = logging.getLogger(__name__)

TEXT_EXTENSIONS = {".txt", ".md"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}


class ExtractionError(Exception):
    """Raised when a document cannot be read at all."""


@dataclass
class ExtractedDocument:
    text: str
    page_count: int = 1
    word_count: int = 0
    char_count: int = 0
    # Signals feeding the ATS rules.
    has_tables: bool = False
    table_count: int = 0
    has_images: bool = False
    image_count: int = 0
    multi_column_pages: int = 0
    has_header_footer_text: bool = False
    is_scanned: bool = False
    chars_per_page: list[int] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    def meta(self) -> dict[str, object]:
        return {
            "page_count": self.page_count,
            "word_count": self.word_count,
            "char_count": self.char_count,
            "has_tables": self.has_tables,
            "table_count": self.table_count,
            "has_images": self.has_images,
            "image_count": self.image_count,
            "multi_column_pages": self.multi_column_pages,
            "has_header_footer_text": self.has_header_footer_text,
            "is_scanned": self.is_scanned,
            "chars_per_page": self.chars_per_page,
            "warnings": self.warnings,
        }


# --------------------------------------------------------------------------- #
# Text hygiene
# --------------------------------------------------------------------------- #
_LIGATURES = {
    "ﬀ": "ff", "ﬁ": "fi", "ﬂ": "fl", "ﬃ": "ffi", "ﬄ": "ffl",
    "‘": "'", "’": "'", "“": '"', "”": '"',
    "–": "-", "—": "-", "−": "-",
    " ": " ", " ": " ", " ": " ", "​": "",
    "": "•", "": "•", "": "•",
}

_BULLET_CHARS = "•●▪◦‣⁃·∙▫■❋❖➜➢→"


def clean_text(raw: str) -> str:
    """Normalise unicode, unify bullets, collapse runaway whitespace."""
    text = unicodedata.normalize("NFKC", raw)
    for src, dst in _LIGATURES.items():
        text = text.replace(src, dst)
    # Strip control chars except tab/newline.
    text = "".join(ch for ch in text if ch == "\n" or ch == "\t" or not unicodedata.category(ch).startswith("C"))
    # Any bullet glyph becomes a leading "• ".
    text = re.sub(rf"^[\s]*[{_BULLET_CHARS}]\s*", "• ", text, flags=re.MULTILINE)
    text = re.sub(rf"[{_BULLET_CHARS}]", "•", text)
    # Trailing spaces, tabs to spaces.
    text = text.replace("\t", "    ")
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    # Never more than one blank line in a row.
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def count_words(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text))


def text_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
def _detect_columns(words: list[dict], page_width: float) -> bool:
    """Heuristic: a page is multi-column when word x-positions cluster into two
    bands separated by a gutter, and both bands carry real text."""
    if len(words) < 40 or not page_width:
        return False
    mid = page_width / 2
    gutter_lo, gutter_hi = mid - page_width * 0.06, mid + page_width * 0.06

    left = right = straddling = 0
    for w in words:
        x0, x1 = w.get("x0", 0.0), w.get("x1", 0.0)
        if x1 < gutter_lo:
            left += 1
        elif x0 > gutter_hi:
            right += 1
        else:
            straddling += 1

    total = left + right + straddling
    if not total:
        return False
    # Both sides substantial, and few words crossing the middle.
    both_sides = min(left, right) / total > 0.22
    clean_gutter = straddling / total < 0.12
    return both_sides and clean_gutter


def _extract_pdf(data: bytes) -> ExtractedDocument:
    import pdfplumber

    pages_text: list[str] = []
    doc = ExtractedDocument(text="")

    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            doc.page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=1.5, y_tolerance=3) or ""
                pages_text.append(page_text)
                doc.chars_per_page.append(len(page_text))

                try:
                    tables = page.find_tables()
                except Exception:  # pdfplumber can fail on odd table geometry
                    tables = []
                if tables:
                    doc.table_count += len(tables)

                images = page.images or []
                doc.image_count += len(images)

                try:
                    words = page.extract_words(x_tolerance=1.5, y_tolerance=3) or []
                except Exception:
                    words = []
                if _detect_columns(words, float(page.width or 0)):
                    doc.multi_column_pages += 1
    except Exception as exc:  # fall back to a simpler reader
        logger.warning("pdfplumber failed (%s); falling back to pypdf", exc)
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            doc.page_count = len(reader.pages)
            pages_text = [(p.extract_text() or "") for p in reader.pages]
            doc.chars_per_page = [len(t) for t in pages_text]
            doc.warnings.append(
                "This PDF needed a fallback parser; some applicant tracking "
                "systems may struggle with it too."
            )
        except Exception as inner:
            raise ExtractionError(
                "This PDF could not be read. It may be corrupt or password-protected."
            ) from inner

    doc.has_tables = doc.table_count > 0
    doc.has_images = doc.image_count > 0
    doc.text = clean_text("\n\n".join(pages_text))

    # An image-heavy PDF with almost no text is a scan or a design export.
    if doc.page_count and len(doc.text) < 200 * doc.page_count:
        doc.is_scanned = True
        doc.warnings.append(
            "Very little selectable text was found. If this resume was exported "
            "as an image or scanned, most parsers will read it as blank."
        )
    return doc


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def _extract_docx(data: bytes) -> ExtractedDocument:
    from docx import Document
    from docx.oxml.ns import qn

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(
            "This .docx file could not be read. If it is an older .doc file, "
            "re-save it as .docx or PDF."
        ) from exc

    doc = ExtractedDocument(text="")
    parts: list[str] = []

    for para in document.paragraphs:
        line = para.text.strip()
        if not line:
            continue
        style = (para.style.name or "").lower() if para.style is not None else ""
        if "list" in style and not line.startswith("•"):
            line = f"• {line}"
        parts.append(line)

    for table in document.tables:
        doc.table_count += 1
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # Deduplicate merged cells that repeat their text.
            deduped: list[str] = []
            for cell in cells:
                if cell and (not deduped or deduped[-1] != cell):
                    deduped.append(cell)
            if deduped:
                parts.append(" | ".join(deduped))

    # Headers and footers: ATS parsers routinely drop these.
    for section in document.sections:
        for container in (section.header, section.footer):
            if container is None:
                continue
            hf_text = "\n".join(p.text.strip() for p in container.paragraphs if p.text.strip())
            if hf_text:
                doc.has_header_footer_text = True
                parts.append(hf_text)

    doc.image_count = len(document.inline_shapes)
    try:
        doc.image_count += len(document.element.body.findall(f".//{qn('w:drawing')}"))
    except Exception:
        pass

    doc.has_tables = doc.table_count > 0
    doc.has_images = doc.image_count > 0
    doc.text = clean_text("\n".join(parts))
    doc.chars_per_page = [len(doc.text)]
    # Word does not expose a reliable page count without rendering; estimate.
    doc.page_count = max(1, round(count_words(doc.text) / 500)) if doc.text else 1
    return doc


# --------------------------------------------------------------------------- #
# Plain text
# --------------------------------------------------------------------------- #
def _extract_text(data: bytes) -> ExtractedDocument:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            raw = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ExtractionError("This text file uses an encoding that could not be decoded.")

    doc = ExtractedDocument(text=clean_text(raw))
    doc.chars_per_page = [len(doc.text)]
    doc.page_count = max(1, round(count_words(doc.text) / 500)) if doc.text else 1
    return doc


# --------------------------------------------------------------------------- #
# Entry point
# --------------------------------------------------------------------------- #
def extract_document(filename: str, data: bytes) -> ExtractedDocument:
    """Extract text and layout signals from an uploaded resume."""
    if not data:
        raise ExtractionError("The uploaded file is empty.")

    suffix = PurePath(filename).suffix.lower()
    if suffix in PDF_EXTENSIONS:
        doc = _extract_pdf(data)
    elif suffix in DOCX_EXTENSIONS:
        doc = _extract_docx(data)
    elif suffix in TEXT_EXTENSIONS:
        doc = _extract_text(data)
    elif suffix == ".doc":
        raise ExtractionError(
            "Legacy .doc files are not supported. Re-save the file as .docx or PDF."
        )
    else:
        raise ExtractionError(
            f"Unsupported file type '{suffix or 'unknown'}'. Upload a PDF, DOCX or TXT file."
        )

    if len(doc.text.strip()) < 100:
        raise ExtractionError(
            "Almost no text could be read from this file. If it is a scan or an "
            "exported image, upload a text-based PDF or a DOCX instead."
        )

    doc.word_count = count_words(doc.text)
    doc.char_count = len(doc.text)
    return doc
